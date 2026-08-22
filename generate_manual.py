import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_devanagari_run(paragraph, text, size=11, color=None, bold=True):
    """Adds Devanagari text with low-level Word XML elements so Word renders script properly without tofu boxes."""
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

    # Low-level XML fix for Word Complex Script (CS) font mapping
    rPr = run._element.get_or_add_rPr()
    
    # 1. Set Complex Script Font Name
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Mangal')
    rFonts.set(qn('w:hAnsi'), 'Mangal')
    rFonts.set(qn('w:cs'), 'Mangal')
    rFonts.set(qn('w:hint'), 'cs')  # Forces Word to use Complex Script engine
    rPr.append(rFonts)

    # 2. Force Complex Script Size & Bold
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size * 2)))
    rPr.append(szCs)

    if bold:
        bCs = OxmlElement('w:bCs')
        rPr.append(bCs)

    return run

def create_manual():
    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Color Palette
    PRIMARY_RED = RGBColor(193, 18, 31)      # #C1121F
    DARK_BLUE = RGBColor(0, 48, 73)         # #003049
    TEXT_DARK = RGBColor(33, 37, 41)         # #212529
    WHITE = RGBColor(255, 255, 255)

    # Set Default Font
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = TEXT_DARK

    # --- TITLE HEADER ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Devanagari Mantra with XML CS Fix
    add_devanagari_run(title_p, "॥ श्री गणेशाय नमः ॥\n", size=14, color=PRIMARY_RED, bold=True)

    run_title = title_p.add_run("Purvanchal Ganeshotsav Mandal\n")
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = DARK_BLUE

    run_sub = title_p.add_run("Receipt & Expense Portal — Official User Manual\n")
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = PRIMARY_RED

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def add_custom_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = DARK_BLUE
        return p

    # --- SECTION 1 ---
    add_custom_heading("1. Getting Started & User Authentication")
    
    p = doc.add_paragraph()
    p.add_run("The portal uses 4-digit PIN authentication to grant access based on representative roles.")
    
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["User Role", "Access Level", "Available Features"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_background(cell, "003049")
        p_cell = cell.paragraphs[0]
        run = p_cell.add_run(h)
        run.font.bold = True
        run.font.color.rgb = WHITE

    data = [
        ("Admin Representative", "Full Privileges", "New Receipt Entry, Expense Management, Analytics, Search & Resend"),
        ("Standard Representative", "View & Search", "Analytics & Stats, Flat-Wise Search, WhatsApp Resend")
    ]

    for row_idx, row_data in enumerate(data, start=1):
        row_cells = table.rows[row_idx].cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            cell = row_cells[col_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            p_cell = cell.paragraphs[0]
            run = p_cell.add_run(text)
            if col_idx == 0:
                run.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- SECTION 2 ---
    add_custom_heading("2. Registering a New Donation Receipt (Admin)")
    
    p1 = doc.add_paragraph(style='List Number')
    r1 = p1.add_run("Navigate to Tab: ")
    r1.font.bold = True
    p1.add_run("Click on the 'New Receipt' tab in the navigation bar.")

    p2 = doc.add_paragraph(style='List Number')
    r2 = p2.add_run("Select Language: ")
    r2.font.bold = True
    p2.add_run("Choose English, ")
    add_devanagari_run(p2, "मराठी", size=11, color=TEXT_DARK, bold=False)
    p2.add_run(", or ")
    add_devanagari_run(p2, "हिंदी", size=11, color=TEXT_DARK, bold=False)
    p2.add_run(" from the top-right language switcher dropdown.")

    steps_receipt = [
        ("Fill Donor Details", "Enter Full Name, valid 10-digit WhatsApp number, and Flat Number (must include exactly one hyphen, e.g., A-608)."),
        ("Enter Amount & Family Count", "Input contribution amount in Rupees (Rs). The Amount in Words auto-fills in the selected language."),
        ("Select Payment Mode", "Choose between Online (UPI) or Cash."),
        ("Save & Share", "Click Save & Generate Receipt. Once saved, click 'Send WhatsApp Receipt Link' to deliver the official receipt image link directly to the donor.")
    ]

    for title, desc in steps_receipt:
        p = doc.add_paragraph(style='List Number')
        r_title = p.add_run(f"{title}: ")
        r_title.font.bold = True
        p.add_run(desc)

    # --- SECTION 3 ---
    add_custom_heading("3. Managing Expenses (Admin)")

    steps_expense = [
        ("Category / Event Selection", "Choose the event (e.g., Ganesh Chaturthi Utsav, 15th August, Common Expenses)."),
        ("Expense Details", "Enter Expense Date, Vendor/Payee Name, Amount (Rs), and Description."),
        ("Payment Mode", "Select whether the expense was paid via Online / Bank or Cash."),
        ("Save & Print Statement", "Click Save Expense Record. Under Category-Wise Expenses Ledger, click 'Print Statement' to generate a clean PDF statement for society audits.")
    ]

    for title, desc in steps_expense:
        p = doc.add_paragraph(style='List Number')
        r_title = p.add_run(f"{title}: ")
        r_title.font.bold = True
        p.add_run(desc)

    # --- SECTION 4 ---
    add_custom_heading("4. Financial Analytics & Available Balance")

    p = doc.add_paragraph("The Analytics & Stats tab provides real-time financial tracking:")
    
    analytics_items = [
        ("Tower-Wise Collections", "Shows total collections, receipt counts, and flat participation progress bars for Buildings A, B, C, D1, D2, E, and F1."),
        ("Available Net Balance", "Calculates real-time funds in hand split by Cash in Hand and Online/Bank balances:\n"
                                   "• Cash in Hand = Total Cash Collected - Total Cash Expenses\n"
                                   "• Online / Bank = Total Online Collected - Total Online Expenses\n"
                                   "• Total Available = Total Collections - Total Expenses"),
        ("Daily Cash Flow Ledger", "Displays date-wise collection, expenses, and running cumulative balance over time.")
    ]

    for title, desc in analytics_items:
        p = doc.add_paragraph(style='List Bullet')
        r_title = p.add_run(f"{title}: ")
        r_title.font.bold = True
        p.add_run(desc)

    # --- CALLOUT BOX ---
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    callout_table = doc.add_table(rows=1, cols=1)
    callout_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_cell = callout_table.rows[0].cells[0]
    set_cell_background(c_cell, "FFFDF5")
    set_cell_margins(c_cell, top=140, bottom=140, left=180, right=180)
    
    # Left Red Border styling
    tcPr = c_cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="single" w:sz="24" w:space="0" w:color="C1121F"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
    tcPr.append(tcBorders)

    cp = c_cell.paragraphs[0]
    c_run_title = cp.add_run("Pro Tip — Resending Receipts:\n")
    c_run_title.font.bold = True
    c_run_title.font.color.rgb = PRIMARY_RED
    cp.add_run("You can quickly search for any flat in the 'Search Receipts' tab and click 'Resend WA' to resend the official WhatsApp message and Drive PDF link anytime!")

    # Save Document
    filename = "Purvanchal_C_Portal_User_Manual.docx"
    doc.save(filename)
    print(f"✓ Document successfully created: {filename}")

if __name__ == "__main__":
    create_manual()